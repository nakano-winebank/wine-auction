#!/usr/bin/env python3
"""提携提案のMarkdownをA4のWord文書（→PDF）に変換する簡易コンバータ。

対応記法: # 見出し / --- 区切り / パイプ表 / > 引用 / - 箇条書き / 1. 番号 / **太字**
提案書1本のために書いた最小限の実装で、汎用のMarkdown処理系ではない。
"""
import re
import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JP_FONT = "游ゴシック"   # Word側の既定。PDF生成時は環境にあるフォントへ差し替える
GOLD = RGBColor(0xA7, 0x84, 0x50)   # WineBankのブランドゴールド
INK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_font(run, size=10.5, bold=False, color=INK, font=None):
    font = font or JP_FONT
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # 東アジア用フォントはrFontsのeastAsiaにも指定しないと反映されない
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_runs(par, text, size=10.5, color=INK, bold_all=False):
    """**...** を太字にしつつテキストを流し込む。"""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        set_font(par.add_run(chunk), size=size, bold=bold_all or bool(i % 2), color=color)


def hr(doc):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    set_font(par.add_run("─" * 46), size=8, color=RGBColor(0xCC, 0xC0, 0xAA))


def build_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, text in enumerate(header):
        cell = table.cell(0, c)
        cell.text = ""
        add_runs(cell.paragraphs[0], text, size=9, bold_all=True)
        shade = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:fill"), "F3EDE1")
        cell._tc.get_or_add_tcPr().append(shade)
    for r, row in enumerate(body, start=1):
        for c in range(len(header)):
            cell = table.cell(r, c)
            cell.text = ""
            add_runs(cell.paragraphs[0], row[c] if c < len(row) else "", size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    for attr, cm in (("top", 2.0), ("bottom", 2.0), ("left", 2.0), ("right", 2.0)):
        setattr(section, f"{attr}_margin", Cm(cm))

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            hr(doc)
            i += 1
            continue

        # 表: ヘッダ行 + 区切り行 + 本体
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = [split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            build_table(doc, rows)
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            par = doc.add_paragraph()
            fmt = par.paragraph_format
            if level == 1:
                fmt.space_before, fmt.space_after = Pt(0), Pt(14)
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(par, text, size=18, color=GOLD, bold_all=True)
            elif level == 2:
                fmt.space_before, fmt.space_after = Pt(14), Pt(6)
                add_runs(par, text, size=13, color=GOLD, bold_all=True)
            else:
                fmt.space_before, fmt.space_after = Pt(10), Pt(4)
                add_runs(par, text, size=11, bold_all=True)
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.6)
            par.paragraph_format.space_before = Pt(6)
            par.paragraph_format.space_after = Pt(6)
            add_runs(par, " ".join(b for b in block if b), size=9.5, color=GRAY)
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.left_indent = Cm(0.8 + 0.5 * (len(m.group(1)) // 2))
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, m.group(2))
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.8)
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, f"{m.group(2)}. {m.group(3)}")
            i += 1
            continue

        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        add_runs(par, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) > 3:
        globals()["JP_FONT"] = sys.argv[3]
    convert(sys.argv[1], sys.argv[2])
